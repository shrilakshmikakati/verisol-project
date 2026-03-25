"""
Algorithm 1: E-Contract Knowledge Graph
Supports .docx, .txt files. Multi-page detection → NLP entity extraction → KG build → PNG render.

Key fixes over previous version:
  1. PARTY min_len lowered to 2 — short names (Bob, AB Corp) no longer dropped
  2. 'between X and Y' pattern captures BOTH named parties
  3. Date nodes store raw value as label (not mangled 'date_DDMMYYYY')
  4. Clause nodes (OBLIGATION etc.) store the FULL SENTENCE not just the keyword
  5. _link_obligations_to_parties uses the trigger keyword for sentence matching,
     not the full clause text (which is too long to substring-match)
  6. Payment dedup merges by numeric amount; richer label wins
  7. _build_node_index indexes 3-letter words (not just 4-letter)
  8. PAYMENT linking distinguishes payer (PAYS) from payee (RECEIVES)
"""
import re, io, base64, warnings
from datetime import datetime as dt
from pathlib import Path
import spacy, networkx as nx
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

warnings.filterwarnings("ignore", category=UserWarning, module="matplotlib")

# ── spaCy loader ───────────────────────────────────────────────────────────────
_nlp = None
def get_nlp():
    global _nlp
    if _nlp: return _nlp
    for m in ("en_core_web_lg", "en_core_web_md", "en_core_web_sm"):
        try: _nlp = spacy.load(m); return _nlp
        except Exception: pass
    import subprocess, sys
    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
    _nlp = spacy.load("en_core_web_sm"); return _nlp

# ── Text extraction ────────────────────────────────────────────────────────────
def extract_text_from_docx(path: str) -> str:
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(path); parts = []
    def _pt(elem):
        return "".join(n.text or "" for n in elem.iter(qn("w:t")))
    def _walk(elem):
        for c in elem:
            loc = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if loc == "p":
                t = _pt(c).strip()
                if t: parts.append(t)
            elif loc == "tbl":
                for tr in c.findall(".//" + qn("w:tr")):
                    row = [" ".join(_pt(p).strip() for p in tc.findall(".//" + qn("w:p"))
                                    if _pt(p).strip())
                           for tc in tr.findall(".//" + qn("w:tc"))]
                    if any(row): parts.append(" | ".join(r for r in row if r))
            else:
                _walk(c)
    _walk(doc.element.body)
    return "\n".join(p for p in parts if p.strip())

def extract_text_from_file(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(path)
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

# ── Multi-page extraction: DOCX ────────────────────────────────────────────────
def extract_pages_from_docx(path: str) -> list:
    try:
        from docx import Document
        from docx.oxml.ns import qn as _qn
    except ImportError:
        return []
    doc = Document(path); paras = doc.paragraphs
    if not paras: return []

    def _build_pages(chunks):
        pages = []
        for lines in chunks:
            content = "\n".join(l for l in lines if l.strip())
            if len(content) < 200: continue
            title = next((l.strip()[:50] for l in lines
                          if l.strip() and any(c.isascii() and c.isalpha() for c in l)),
                         f"Page {len(pages)+1}")
            pages.append((len(pages)+1, content, title))
        return pages

    s_chunks, cur = [], []
    for para in paras:
        if para.text.strip(): cur.append(para.text)
        pPr = para._element.find(_qn("w:pPr"))
        if pPr is not None and pPr.find(_qn("w:sectPr")) is not None:
            s_chunks.append(cur[:]); cur = []
    if cur: s_chunks.append(cur)
    pages = _build_pages(s_chunks)
    if len(pages) >= 2: return pages

    pb_chunks, cur = [], []
    for para in paras:
        xml = para._element.xml if hasattr(para._element, "xml") else ""
        if ("lastRenderedPageBreak" in xml or 'w:type="page"' in xml) and cur:
            pb_chunks.append(cur[:]); cur = []
        if para.text.strip(): cur.append(para.text)
    if cur: pb_chunks.append(cur)
    pages = _build_pages(pb_chunks)
    if len(pages) >= 2: return pages

    heading_markers = [(i, para.text.strip())
                       for i, para in enumerate(paras)
                       if para.style and any(h in para.style.name for h in
                          ("Heading 1","Heading 2","Heading 3")) and para.text.strip()]
    if len({t for _, t in heading_markers}) >= 2 or len(heading_markers) >= 3:
        h_chunks = []
        for j, (start, _) in enumerate(heading_markers):
            end = heading_markers[j+1][0] if j+1 < len(heading_markers) else len(paras)
            h_chunks.append([paras[k].text for k in range(start, end) if paras[k].text.strip()])
        pages = _build_pages(h_chunks)
        if len(pages) >= 2: return pages

    total = sum(len(p.text) for p in paras)
    cpc = 4000 if total > 16000 else 3000 if total > 8000 else 2000 if total > 4000 else 1500
    ac_chunks, cur, count = [], [], 0
    for para in paras:
        t = para.text.strip()
        if not t: continue
        cur.append(para.text); count += len(t)
        if count >= cpc:
            ac_chunks.append(cur[:]); cur = []; count = 0
    if cur: ac_chunks.append(cur)
    pages = _build_pages(ac_chunks)
    return pages if len(pages) >= 2 else []

# ── Multi-page extraction: TXT ─────────────────────────────────────────────────
def extract_pages_from_txt(path: str) -> list:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
    except OSError:
        return []
    if not raw.strip(): return []

    def _build(chunks):
        pages = []
        for chunk in chunks:
            content = chunk.strip()
            if len(content) < 200: continue
            first = next((l.strip() for l in content.splitlines() if l.strip()), "")
            title = first[:50] if first else f"Page {len(pages)+1}"
            pages.append((len(pages)+1, content, title))
        return pages

    if "\f" in raw:
        pages = _build(raw.split("\f"))
        if len(pages) >= 2: return pages

    blocks = re.split(r"\n{3,}", raw)
    if len(blocks) >= 2:
        heading_re = re.compile(
            r"^(?:[A-Z][A-Z\s\-:,\.]{3,}|[A-Z][a-z][\w\s\-:,\.]{3,}|\d+[\.\-\)]\s+\S)")
        starts = [i for i, b in enumerate(blocks)
                  if heading_re.match(next((l.strip() for l in b.splitlines() if l.strip()), ""))]
        if len(starts) >= 2:
            chunks = ["\n\n".join(blocks[s: starts[j+1] if j+1 < len(starts) else len(blocks)])
                      for j, s in enumerate(starts)]
            pages = _build(chunks)
            if len(pages) >= 2: return pages

    total = len(raw)
    cpc = 4000 if total > 16000 else 3000 if total > 8000 else 2000 if total > 4000 else 1500
    lines = raw.splitlines()
    chunks, cur, count = [], [], 0
    for line in lines:
        cur.append(line); count += len(line)
        if count >= cpc:
            chunks.append("\n".join(cur)); cur = []; count = 0
    if cur: chunks.append("\n".join(cur))
    pages = _build(chunks)
    return pages if len(pages) >= 2 else []

# ── Unified entry point ────────────────────────────────────────────────────────
def extract_pages_from_file(path: str, *, fallback_single: bool = False) -> list:
    p = Path(path); ext = p.suffix.lower()
    if ext == ".docx": pages = extract_pages_from_docx(path)
    elif ext == ".txt": pages = extract_pages_from_txt(path)
    else: return []
    if pages: return pages
    if fallback_single: return _single_page_fallback(path)
    return []

def _single_page_fallback(path: str) -> list:
    try: text = extract_text_from_file(path).strip()
    except Exception: return []
    if not text: return []
    first = next((l.strip() for l in text.splitlines() if l.strip()), "Document")
    return [(1, text, first[:60])]

# ── Preprocessing ──────────────────────────────────────────────────────────────
def preprocess_text(text: str) -> str:
    text = re.sub(r'[\u0900-\u097F\u4e00-\u9fff]+', ' ', text)
    for _ in range(4):
        prev = text
        text = re.sub(r'(?:\b\d{6,}\b\s*){2,}', ' ', text)
        if text == prev: break
    text = re.sub(r'(?<!\d)\d{6}(?![\d/\-])', '', text)
    text = text.replace("\u2019","'").replace("\u2018","'")
    text = text.replace("\u201c",'"').replace("\u201d",'"')
    text = text.replace("\u2013","-").replace("\u2014"," - ")
    text = re.sub(r'&#?\w+;', '', text)
    return re.sub(r'\s+', ' ', text).strip()

def _ascii_label(s: str) -> str:
    clean = re.sub(r"[^\x00-\x7F]+", "?", s)
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean[:30] if clean.replace("?","").strip() else s[:20]

def detect_language(text: str) -> str:
    sample = text[:500].lower()
    scores = {
        "en": len(re.findall(r"\b(?:the|shall|will|party|agreement|contract)\b", sample)),
        "de": len(re.findall(r"\b(?:der|die|das|vertrag|partei|soll|muss)\b", sample)),
        "fr": len(re.findall(r"\b(?:le|la|les|contrat|partie|doit|accord)\b", sample)),
        "es": len(re.findall(r"\b(?:el|la|los|contrato|parte|debe|acuerdo)\b", sample)),
        "hi": len(re.findall(r"[\u0900-\u097F]", sample)),
        "zh": len(re.findall(r"[\u4e00-\u9fff]", sample)),
    }
    return max(scores, key=scores.get)

# ── Entity Blocklist ───────────────────────────────────────────────────────────
ENTITY_BLOCKLIST = {
    "mobile","phone","fax","tel","sir","dear","sub","madam",
    "regards","yours","sincerely","faithfully","truly",
    "the","and","for","from","to","in","of","on","at","by","is","be",
    "as","an","or","if","re","a","it","its","this","that","these",
    "etc","viz","ie","eg","nb","pp","cf","vs","op",
    "govt","estt","sno","sl","sr","advt","circular","memo","minutes",
}

def _is_valid_entity(val: str, etype: str) -> bool:
    v = val.strip(); vl = v.lower()
    # Permissive minimums — 2 chars for party names, 3 for everything else
    min_len = 2 if etype == "PARTY" else 3
    if len(v) < min_len: return False
    if vl in ENTITY_BLOCKLIST: return False
    if etype == "PARTY":
        if re.fullmatch(r"[\d\s\+\-\(\)\.]+", v): return False
        if re.fullmatch(r"\d{5,}", v.replace(" ","")): return False
        # Only reject very short pure-lowercase words (1-2 chars like "is", "at")
        if re.fullmatch(r"[a-z]{1,2}", v): return False
    if etype == "PAYMENT" and not re.search(r"\d", v): return False
    if etype == "DATE_DEADLINE":
        if not re.search(r"\d", v):
            if not re.search(
                r"\b(?:effective|commencement|expiry|joining|reporting|start|end)\s+date\b",
                v, re.I): return False
    return True

# ── Pattern Library ────────────────────────────────────────────────────────────
PARTY_PATTERNS = [
    # Role titles
    r"\b(Buyer|Seller|Vendor|Client|Customer|Contractor|Subcontractor|"
    r"Licensor|Licensee|Franchisor|Franchisee|Lessor|Lessee|Landlord|Tenant|"
    r"Employer|Employee|Principal|Agent|Borrower|Lender|Creditor|Debtor|"
    r"Guarantor|Surety|Service\s+Provider|Service\s+Recipient|"
    r"Intern|Candidate|Supervisor|Director|Professor|Faculty|"
    r"Researcher|Scientist|Fellow|Scholar|Trainee|Apprentice|"
    r"Owner|Developer|Consultant|Advisor|Partner)\b",
    # Indian institution acronyms
    r"\b((?:NIT|IIT|IIM|AIIMS|BITS|TIFR|ISRO|DRDO|CSIR|BARC)\s+[A-Za-z]+(?:\s+[A-Za-z]+)?)\b",
    # Company names with legal suffixes
    r"\b([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,4})\s+"
    r"(?:Inc\.|LLC|Ltd\.|Corp\.|GmbH|S\.A\.|B\.V\.|PLC|LLP|LP|AG|SA)\b",
    # Titled persons
    r"\b(?:Dr\.|Mr\.|Ms\.|Mrs\.|Prof\.)\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,3})\b",
    # Quoted party names like ("the Company") or ('Party A')
    r'[("\']((?:the\s+)?[A-Z][a-zA-Z\s]{2,30})[)"\']',
    # "between X and Y" — captures BOTH parties
    r"\bbetween\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,3})\s+(?:\(.*?\)\s+)?and\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,3})\b",
    # "X (the Landlord/Tenant/etc.)" — extracts proper name
    r"\b([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,2})\s+\(the\s+"
    r"(?:Landlord|Tenant|Employer|Employee|Buyer|Seller|Vendor|Client|"
    r"Licensor|Licensee|Lessor|Lessee|Borrower|Lender|Owner|Contractor|"
    r"Intern|Supervisor|Director|Principal|Agent)\)",
]

PAYMENT_PATTERNS = [
    # Currency prefix (Rs, INR, USD, $, €, £) + amount
    r"(?:Rs\.?\s*|INR\s*|USD\s*|\$|€|£)[\d,]+(?:\.\d{1,2})?"
    r"(?:\s*(?:/-\s*)?(?:per\s+month|p\.m\.|pm|/month|/annum|lakhs?|crores?|thousands?))?\b",
    # "X per month/annum" without currency prefix
    r"\b[\d,]+(?:\.\d{1,2})?\s*(?:per\s+month|p\.m\.|/month|per\s+annum)\b",
    # "stipend/salary/rent/fee/deposit of [Rs.] amount"
    r"\b(?:stipend|salary|remuneration|honorarium|allowance|fellowship|"
    r"emolument|wages?|pay|rent|fee|deposit|premium|consideration|amount)\s+of\s+"
    r"(?:Rs\.?\s*|INR\s*|USD\s*|\$|€|£)?[\d,]+(?:\.\d{1,2})?\b",
    r"\bconsolidated\s+(?:stipend|salary)\s+of\s+(?:Rs\.?\s*|INR\s*)?[\d,]+\b",
    r"Rs\.?\s*[\d,]+\s*(?:/-\s*)?(?:per\s+month|p\.m\.|/month)",
    # Security deposit, late fee, penalty — common in rental/service contracts
    r"\b(?:security\s+deposit|late\s+fee|penalty|fine|liquidated\s+damages)\s+of\s+"
    r"(?:Rs\.?\s*|INR\s*|\$|€|£)?[\d,]+(?:\.\d{1,2})?\b",
    # Bare currency amounts (last resort, short pattern)
    r"(?:Rs\.?\s*|INR\s*|\$|€|£)[\d,]+(?:\.\d{1,2})?",
]

DATE_PATTERNS = [
    r"\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b",
    r"\b\d{4}-\d{2}-\d{2}\b",
    r"(?:0[1-9]|[12]\d|3[01])(?:0[1-9]|1[0-2])(?:19|20)\d{2}",
    r"\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
    r"\b(?:January|February|March|April|May|June|July|August|"
    r"September|October|November|December)\s+\d{1,2},?\s+\d{4}\b",
    r"\b\d+\s+(?:calendar\s+)?(?:days?|business\s+days?|weeks?|months?|years?)\b",
    r"\b(?:effective|commencement|expiry|joining|reporting|start|end|before|after)\s+date\b",
]
DISPUTE_PATTERNS      = [r"\b(?:arbitration|mediation|dispute\s+resolution|governing\s+law|"
                          r"jurisdiction|competent\s+court|ICC|AAA|UNCITRAL|LCIA|SIAC)\b"]
CONFIDENTIAL_PATTERNS = [r"\b(?:confidential(?:ity)?|non[\s\-]?disclosure|proprietary\s+information|"
                          r"trade\s+secret|intellectual\s+property|IP\s+rights?|copyright|patent|trademark)\b"]
FORCE_MAJEURE_PATTERNS= [r"\b(?:force\s+majeure|act\s+of\s+God|beyond\s+(?:the\s+)?"
                          r"(?:reasonable\s+)?control|pandemic|natural\s+disaster|"
                          r"government\s+(?:action|order))\b"]
MILESTONE_PATTERNS    = [r"\b(?:milestone|deliverable|phase\s+[1-9]|stage\s+[1-9]|"
                          r"checkpoint|acceptance\s+criteria|sign[\s\-]?off|handover)\b"]

# Trigger keywords — used for sentence matching; clause nodes store the FULL sentence
OBLIGATION_TRIGGERS  = ["shall","must","is required to","agrees to","undertakes to","covenants to",
    "is obligated to","is bound to","will provide","shall deliver","shall supply","shall report",
    "shall complete","shall maintain","shall submit","shall perform","shall pay","shall receive",
    "will receive","shall be paid","must report","reasonable efforts","best efforts"]
CONDITION_TRIGGERS   = ["if and only if","provided that","subject to","on condition that",
    "contingent upon","in the event that","unless","notwithstanding","except where",
    "save as otherwise","to the extent that","condition precedent","failing which","in case of"]
TERMINATION_TRIGGERS = ["terminate","termination","cancellation","rescind","void","expire","expiry",
    "cancelled","treat as cancelled","terminable","terminated at any time",
    "appointment is cancelled","appointment is purely temporary"]
PENALTY_TRIGGERS     = ["penalty","liquidated damages","late fee","interest per",
    "default interest","indemnify","indemnification","hold harmless"]

# ── Date normalisation ─────────────────────────────────────────────────────────
def normalize_date(date_str: str) -> str:
    """Normalise any date string to DDMMYYYY."""
    s = date_str.strip()
    if re.match(r"^\d{8}$", s):
        for fmt in ("%d%m%Y", "%Y%m%d"):
            try: return dt.strptime(s, fmt).strftime("%d%m%Y")
            except ValueError: pass
    if re.match(r"^\d{1,2}[/-]\d{1,2}[/-]\d{4}$", s):
        try: return dt.strptime(s.replace("-","/"), "%d/%m/%Y").strftime("%d%m%Y")
        except ValueError: pass
    if re.match(r"^\d{4}-\d{2}-\d{2}$", s):
        try: return dt.strptime(s, "%Y-%m-%d").strftime("%d%m%Y")
        except ValueError: pass
    for fmt in ("%d %B %Y","%d %b %Y","%B %d, %Y","%b %d, %Y","%B %d %Y","%b %d %Y"):
        try: return dt.strptime(s, fmt).strftime("%d%m%Y")
        except ValueError: pass
    return s

def _date_to_unix(normalized_ddmmyyyy: str) -> int:
    """Convert DDMMYYYY string to Unix timestamp. Returns 0 on failure."""
    s = normalized_ddmmyyyy.strip()
    if len(s) == 8 and s.isdigit():
        try:
            dd, mm, yyyy = int(s[:2]), int(s[2:4]), int(s[4:])
            if 1 <= mm <= 12 and 1 <= dd <= 31 and 1970 <= yyyy <= 2100:
                return int(dt(yyyy, mm, dd).timestamp())
        except Exception:
            pass
    return 0

# ── spaCy entity extraction ────────────────────────────────────────────────────
def extract_entities_spacy(text: str) -> list:
    nlp = get_nlp(); doc = nlp(text[:80000])
    LABEL_MAP = {"PERSON":"PARTY","ORG":"PARTY","GPE":"PARTY",
                 "DATE":"DATE_DEADLINE","TIME":"DATE_DEADLINE",
                 "MONEY":"PAYMENT","LAW":"CONFIDENTIALITY_IP"}
    entities, seen = [], set()
    for ent in doc.ents:
        mapped = LABEL_MAP.get(ent.label_)
        if not mapped: continue
        val = ent.text.strip()
        if not _is_valid_entity(val, mapped): continue
        if len(val) > 80: continue
        if re.search(r"\b(?:shall|must|will|if|unless|because|therefore)\b", val, re.I): continue
        key = (mapped, val.lower()[:50])
        if key not in seen:
            seen.add(key); entities.append({"type": mapped, "value": val, "raw": val})
    return entities

def extract_entities_atomic(text: str) -> list:
    """
    Pattern-based extraction. Raw match value is preserved exactly.
    'between X and Y' pattern captures both party names.
    """
    entities, seen = [], set()

    def add(etype: str, raw: str):
        val = re.sub(r"\s+", " ", raw).strip()[:80]
        if not _is_valid_entity(val, etype): return
        key = (etype, val.lower()[:60])
        if key not in seen:
            seen.add(key); entities.append({"type": etype, "value": val, "raw": val})

    for pat in PARTY_PATTERNS:
        for m in re.finditer(pat, text):
            if m.lastindex and m.lastindex >= 2:
                # "between X and Y" — two groups
                add("PARTY", m.group(1))
                add("PARTY", m.group(2))
            elif m.lastindex:
                add("PARTY", m.group(1))
            else:
                add("PARTY", m.group())

    for pat in PAYMENT_PATTERNS:
        for m in re.finditer(pat, text, re.IGNORECASE): add("PAYMENT", m.group())

    for pat in DATE_PATTERNS:
        for m in re.finditer(pat, text, re.IGNORECASE): add("DATE_DEADLINE", m.group())

    for etype, pats in [("DISPUTE_ARBITRATION", DISPUTE_PATTERNS),
                        ("CONFIDENTIALITY_IP",  CONFIDENTIAL_PATTERNS),
                        ("FORCE_MAJEURE",       FORCE_MAJEURE_PATTERNS),
                        ("MILESTONE",           MILESTONE_PATTERNS)]:
        for pat in pats:
            for m in re.finditer(pat, text, re.IGNORECASE): add(etype, m.group())

    return entities

def extract_clause_entities(text: str) -> list:
    """
    Extract OBLIGATION/CONDITION/TERMINATION/PENALTY entities as FULL SENTENCES.

    Critical fix: the old code stored only the keyword ('shall') making all
    OBLIGATION nodes useless stubs. Now we store the complete clause sentence
    so business logic (who does what, to whom, by when, for how much) is
    preserved and carried into Solidity generation.
    """
    nlp = get_nlp(); doc = nlp(text[:60000])
    kw_map = [
        ("OBLIGATION",    OBLIGATION_TRIGGERS),
        ("CONDITION",     CONDITION_TRIGGERS),
        ("TERMINATION",   TERMINATION_TRIGGERS),
        ("PENALTY_REMEDY", PENALTY_TRIGGERS),
    ]
    entities, seen = [], set()
    for sent in doc.sents:
        s_lower = sent.text.lower()
        s_text  = sent.text.strip()
        if len(s_text) < 10: continue
        for etype, triggers in kw_map:
            for kw in triggers:
                if kw in s_lower:
                    clause_val = re.sub(r"\s+", " ", s_text)[:120]
                    key = (etype, clause_val.lower()[:80])
                    if key not in seen:
                        seen.add(key)
                        entities.append({
                            "type":    etype,
                            "value":   clause_val,
                            "raw":     s_text,
                            "trigger": kw,
                        })
                    break  # one type classification per sentence

    return entities

# ── Semantic edge extraction ───────────────────────────────────────────────────
SEMANTIC_VERB_MAP = {
    "pay":"PAYS","receive":"RECEIVES","report":"REPORTS_TO","provide":"PROVIDES",
    "deliver":"DELIVERS","terminate":"TERMINATES","assign":"ASSIGNS","complete":"COMPLETES",
    "fulfill":"FULFILLS","submit":"SUBMITS","notify":"NOTIFIES","appoint":"APPOINTS",
    "engage":"ENGAGES","employ":"EMPLOYS","sign":"SIGNS","cancel":"CANCELS",
    "breach":"BREACHES","penalize":"PENALIZES","join":"JOINS","work":"WORKS_FOR",
    "supervise":"SUPERVISES","guide":"GUIDES","apply":"APPLIES","require":"REQUIRES",
    "owe":"OWES","deposit":"DEPOSITS","forfeit":"FORFEITS","indemnify":"INDEMNIFIES",
    "grant":"GRANTS","maintain":"MAINTAINS","use":"USES","occupy":"OCCUPIES",
}

def _build_node_index(G: nx.DiGraph) -> dict:
    """Build fuzzy lookup index: normalised_string → node_id."""
    idx = {}
    for n in G.nodes:
        label = G.nodes[n].get("label", n)
        normed = re.sub(r"[^a-z0-9]", "", label.lower())
        idx[normed] = n
        # Index words 3+ chars (not just 4+) to catch short names like "Bob"
        for word in re.findall(r"[a-z]{3,}", label.lower()):
            if word not in idx: idx[word] = n
    return idx

def _fuzzy_node_lookup(phrase: str, node_idx: dict):
    if not phrase or len(phrase) < 2: return None
    normed = re.sub(r"[^a-z0-9]", "", phrase.lower())
    if normed in node_idx: return node_idx[normed]
    for w in re.findall(r"[a-z]{3,}", phrase.lower()):
        if w in node_idx: return node_idx[w]
    for key in node_idx:
        if len(normed) >= 3 and (normed.startswith(key) or key.startswith(normed)):
            return node_idx[key]
    return None

def extract_semantic_edges(text: str, G: nx.DiGraph):
    """Extract subject→verb→object triples; store relation + source sentence on each edge."""
    nlp = get_nlp(); node_idx = _build_node_index(G)
    for sent in nlp(text[:40000]).sents:
        tokens = list(sent)
        subjs  = [t for t in tokens if t.dep_ in ("nsubj","nsubjpass")]
        objs   = [t for t in tokens if t.dep_ in ("dobj","pobj","attr")]
        verbs  = [t for t in tokens if t.pos_ == "VERB"]
        for subj in subjs:
            subj_phrase = " ".join(t.text for t in subj.subtree
                                   if t.pos_ in ("NOUN","PROPN","ADJ") or t == subj)
            src = _fuzzy_node_lookup(subj_phrase, node_idx) or _fuzzy_node_lookup(subj.text, node_idx)
            if not src: continue
            verb = min(verbs, key=lambda v: abs(v.i - subj.i), default=None)
            if not verb: continue
            lemma = verb.lemma_.lower()
            neg = any(c.dep_ == "neg" for c in verb.children)
            sem_label = SEMANTIC_VERB_MAP.get(lemma)
            if not sem_label: continue
            if neg: sem_label = "NOT_" + sem_label
            for obj in objs:
                obj_phrase = " ".join(t.text for t in obj.subtree
                                      if t.pos_ in ("NOUN","PROPN","ADJ") or t == obj)
                tgt = _fuzzy_node_lookup(obj_phrase, node_idx) or _fuzzy_node_lookup(obj.text, node_idx)
                if tgt and src != tgt and not G.has_edge(src, tgt):
                    G.add_edge(src, tgt,
                               relation=sem_label,
                               sentence=sent.text.strip()[:120])

# ── Obligation → Party linking ─────────────────────────────────────────────────
def _link_obligations_to_parties(text: str, G: nx.DiGraph):
    """
    For each clause node find the co-occurring party and add HAS_OBLIGATION edge.
    For payment nodes distinguish payer (PAYS) vs payee (RECEIVES).
    Uses TRIGGER KEYWORD for sentence matching (not full clause text).
    """
    party_nodes   = [n for n in G.nodes if G.nodes[n].get("entity_type") == "PARTY"]
    payment_nodes = [n for n in G.nodes if G.nodes[n].get("entity_type") == "PAYMENT"]
    clause_types  = {"OBLIGATION","CONDITION","TERMINATION","PENALTY_REMEDY"}
    target_nodes  = [n for n in G.nodes if G.nodes[n].get("entity_type") in clause_types]
    sentences     = re.split(r"[.;\n]", text)

    for clause_node in target_nodes:
        trigger = G.nodes[clause_node].get("trigger", "").lower()
        if not trigger: continue
        counts = {}
        for sent in sentences:
            sl = sent.lower()
            if trigger not in sl: continue
            for p in party_nodes:
                plabel = G.nodes[p].get("label", p).lower()
                words  = re.findall(r"[a-z]{3,}", plabel)
                if plabel in sl or any(w in sl for w in words):
                    counts[p] = counts.get(p, 0) + 1
        if counts:
            best = max(counts, key=counts.get)
            if not G.has_edge(best, clause_node):
                G.add_edge(best, clause_node,
                           relation="HAS_OBLIGATION",
                           sentence=G.nodes[clause_node].get("raw","")[:80])

    pay_verbs = re.compile(
        r"\b(?:pay|pays|paid|receive|receives|received|deposit|deposits|"
        r"stipend|salary|rent|fee|amount|consideration)\b", re.I)
    for pay_node in payment_nodes:
        pay_label  = G.nodes[pay_node].get("label", pay_node)
        pay_amount = re.search(r"\d[\d,]*", pay_label)
        for sent in sentences:
            sl = sent.lower()
            if not pay_verbs.search(sl): continue
            if pay_amount and pay_amount.group().replace(",","") not in sl.replace(",",""): continue
            for p in party_nodes:
                plabel = G.nodes[p].get("label", p).lower()
                words  = re.findall(r"[a-z]{3,}", plabel)
                if plabel in sl or any(w in sl for w in words):
                    is_payer = bool(re.search(r"\b(?:pay|pays|paid|deposit)\b", sl))
                    rel = "PAYS" if is_payer else "RECEIVES"
                    if not G.has_edge(p, pay_node):
                        G.add_edge(p, pay_node,
                                   relation=rel,
                                   sentence=sent.strip()[:120])

# ── Build KG ───────────────────────────────────────────────────────────────────
def build_econtract_knowledge_graph(raw_text: str) -> nx.DiGraph:
    """
    Full pipeline: preprocess → extract entities → deduplicate → build graph → add edges.

    Node ID: stable string "{ETYPE}::{raw_value[:80]}"
    Node attributes: entity_type, label (=raw value), raw, lang,
                     normalized + unix_ts (date nodes), amount (payment nodes),
                     trigger (clause nodes)
    Edge attributes: relation, sentence (source clause for traceability)
    """
    text = preprocess_text(raw_text)
    lang = detect_language(text)
    G    = nx.DiGraph()
    G.graph["language"] = lang
    G.graph["source_len"] = len(text)

    raw_ents = (extract_entities_atomic(text) +
                extract_entities_spacy(text)  +
                extract_clause_entities(text))

    seen_nids:    set  = set()
    seen_amounts: dict = {}   # normalised_digits → node_id

    def _norm_amount(val: str) -> str:
        m = re.search(r"\d[\d,]*(?:\.\d+)?", val)
        return m.group().replace(",", "") if m else ""

    def _nid(val: str, etype: str) -> str:
        return f"{etype}::{val[:80]}"

    for e in raw_ents:
        val   = e["value"].strip()
        etype = e["type"]
        if not val: continue

        if etype == "DATE_DEADLINE":
            normalized = normalize_date(val)
            unix_ts    = _date_to_unix(normalized)
            is_duration = bool(re.search(r"\d+\s+(?:days?|months?|years?|weeks?)", val, re.I))
            nid = _nid(val, "DATE")
            if nid in seen_nids: continue
            seen_nids.add(nid)
            G.add_node(nid,
                       entity_type="DATE_DEADLINE",
                       label=val,              # ← raw value preserved (e.g. "2023-01-01")
                       normalized=normalized,
                       unix_ts=unix_ts,
                       is_duration=is_duration,
                       raw=e.get("raw", val),
                       lang=lang)

        elif etype == "PAYMENT":
            norm_amt = _norm_amount(val)
            if norm_amt and norm_amt in seen_amounts:
                existing = seen_amounts[norm_amt]
                # Prefer the richer label (longer = more context)
                if len(val) > len(G.nodes[existing].get("label", "")):
                    G.nodes[existing]["label"] = val
                    G.nodes[existing]["raw"]   = e.get("raw", val)
                continue
            nid = _nid(val, "PAYMENT")
            if nid in seen_nids: continue
            seen_nids.add(nid)
            if norm_amt: seen_amounts[norm_amt] = nid
            G.add_node(nid,
                       entity_type="PAYMENT",
                       label=val,
                       amount=norm_amt,
                       raw=e.get("raw", val),
                       lang=lang)

        elif etype in ("OBLIGATION","CONDITION","TERMINATION","PENALTY_REMEDY"):
            nid = _nid(val, etype)
            if nid in seen_nids: continue
            seen_nids.add(nid)
            G.add_node(nid,
                       entity_type=etype,
                       label=val,              # ← full sentence
                       raw=e.get("raw", val),
                       trigger=e.get("trigger", ""),
                       lang=lang)

        else:
            nid = _nid(val, etype)
            if nid in seen_nids: continue
            seen_nids.add(nid)
            G.add_node(nid,
                       entity_type=etype,
                       label=val,
                       raw=e.get("raw", val),
                       lang=lang)

    extract_semantic_edges(text, G)
    _link_obligations_to_parties(text, G)
    return G

def graph_to_dict(G: nx.DiGraph) -> dict:
    return {
        "nodes": [{"id": n, **G.nodes[n]} for n in G.nodes],
        "edges": [{"source": u, "target": v, **G.edges[u,v]} for u, v in G.edges],
        "meta":  {"language": G.graph.get("language","en"),
                  "source_len": G.graph.get("source_len", 0)},
    }

# ── Visualisation ──────────────────────────────────────────────────────────────
TYPE_COLORS = {
    "PARTY":"#38bdf8","OBLIGATION":"#fb923c","DATE_DEADLINE":"#4ade80",
    "PAYMENT":"#facc15","CONDITION":"#c084fc","TERMINATION":"#f87171",
    "PENALTY_REMEDY":"#ff6b6b","DISPUTE_ARBITRATION":"#e879f9",
    "CONFIDENTIALITY_IP":"#a3e635","FORCE_MAJEURE":"#67e8f9",
    "MILESTONE":"#fbbf24","GENERIC":"#64748b",
}

def render_graph_base64(G: nx.DiGraph, title: str = "Knowledge Graph") -> str:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        fig, ax = plt.subplots(figsize=(16, 10))
        ax.set_facecolor("#0f172a"); fig.patch.set_facecolor("#0f172a")
        if G.number_of_nodes() == 0:
            ax.text(0.5, 0.5, "No entities extracted", ha="center",
                    color="white", transform=ax.transAxes)
        else:
            color_map  = [TYPE_COLORS.get(G.nodes[n].get("entity_type","GENERIC"),"#64748b")
                          for n in G.nodes]
            k_val      = max(1.5, 5.0 / max(1, G.number_of_nodes() ** 0.4))
            pos        = nx.spring_layout(G, k=k_val, seed=42)
            node_sizes = [800 if G.nodes[n].get("entity_type") != "GENERIC" else 350
                          for n in G.nodes]
            nx.draw_networkx_nodes(G, pos, node_color=color_map,
                                   node_size=node_sizes, alpha=0.92, ax=ax)
            labels = {n: _ascii_label(G.nodes[n].get("label", n)) for n in G.nodes}
            nx.draw_networkx_labels(G, pos, labels, font_size=6.5, font_color="white", ax=ax)
            nx.draw_networkx_edges(G, pos, edge_color="#334155", arrows=True,
                                   arrowsize=12, connectionstyle="arc3,rad=0.08", ax=ax)
            el = {e: G.edges[e].get("relation","")[:16] for e in G.edges}
            nx.draw_networkx_edge_labels(G, pos, el, font_size=5.5,
                                         font_color="#94a3b8", ax=ax)
            legend = [mpatches.Patch(color=c, label=t) for t, c in TYPE_COLORS.items()]
            ax.legend(handles=legend, loc="upper left", fontsize=6.5,
                      facecolor="#1e293b", labelcolor="white", framealpha=0.85, ncol=2)
            ax.set_title(f"{title}  [lang={G.graph.get('language','en')}]",
                         color="white", fontsize=13, pad=10)
        ax.axis("off")
        buf = io.BytesIO()
        plt.savefig(buf, format="png", bbox_inches="tight", dpi=120,
                    facecolor=fig.get_facecolor())
        plt.close(fig)
    return base64.b64encode(buf.getvalue()).decode()